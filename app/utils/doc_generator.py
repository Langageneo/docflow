"""
Générateur de documents Excel, PDF, Word
"""

import io
import pandas as pd
import numpy as np
from datetime import datetime


def generate_excel_report(df, title="Rapport", company="DocFlow", author="DocFlow", lang="fr"):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Données")
        workbook = writer.book
        worksheet = writer.sheets["Données"]

        from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
        from openpyxl.utils import get_column_letter

        header_fill = PatternFill(start_color="E94560", end_color="E94560", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=11)
        alt_fill = PatternFill(start_color="1A1A2E", end_color="1A1A2E", fill_type="solid")
        alt_fill2 = PatternFill(start_color="16213E", end_color="16213E", fill_type="solid")
        normal_font = Font(color="EAEAEA", size=10)
        thin_border = Border(
            bottom=Side(style="thin", color="2A2A4A"),
            right=Side(style="thin", color="2A2A4A"),
        )

        for col_idx, col_name in enumerate(df.columns, start=1):
            cell = worksheet.cell(row=1, column=col_idx)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = thin_border

        for row_idx in range(2, len(df) + 2):
            fill = alt_fill if row_idx % 2 == 0 else alt_fill2
            for col_idx in range(1, len(df.columns) + 1):
                cell = worksheet.cell(row=row_idx, column=col_idx)
                cell.fill = fill
                cell.font = normal_font
                cell.alignment = Alignment(horizontal="left", vertical="center")
                cell.border = thin_border

        for col_idx, col_name in enumerate(df.columns, start=1):
            col_letter = get_column_letter(col_idx)
            max_length = max(len(str(col_name)), df[col_name].astype(str).str.len().max() if len(df) > 0 else 0)
            worksheet.column_dimensions[col_letter].width = min(max_length + 4, 40)

        worksheet.row_dimensions[1].height = 25
        worksheet.freeze_panes = "A2"

        numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
        if numeric_cols:
            stats_data = df[numeric_cols].describe().reset_index()
            stats_sheet = workbook.create_sheet(title="Statistiques" if lang == "fr" else "Statistics")
            for col_idx, col_name in enumerate(stats_data.columns, start=1):
                cell = stats_sheet.cell(row=1, column=col_idx, value=str(col_name))
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal="center")
            for row_idx, row in enumerate(stats_data.itertuples(index=False), start=2):
                for col_idx, val in enumerate(row, start=1):
                    cell = stats_sheet.cell(row=row_idx, column=col_idx,
                                            value=round(float(val), 3) if isinstance(val, float) else val)
                    cell.font = Font(color="EAEAEA", size=10)
                    cell.fill = alt_fill if row_idx % 2 == 0 else alt_fill2

        summary_sheet = workbook.create_sheet(title="Résumé" if lang == "fr" else "Summary")
        summary_data = [
            ["DocFlow — Rapport", title],
            ["Date", datetime.now().strftime("%d/%m/%Y %H:%M")],
            ["Société", company],
            ["Auteur", author],
            ["Lignes" if lang == "fr" else "Rows", len(df)],
            ["Colonnes" if lang == "fr" else "Columns", len(df.columns)],
            ["Valeurs manquantes" if lang == "fr" else "Missing Values", int(df.isnull().sum().sum())],
        ]
        for row_idx, (key, val) in enumerate(summary_data, start=1):
            summary_sheet.cell(row=row_idx, column=1, value=key).font = Font(bold=True, color="E94560")
            summary_sheet.cell(row=row_idx, column=2, value=val).font = Font(color="EAEAEA")
        summary_sheet.column_dimensions["A"].width = 30
        summary_sheet.column_dimensions["B"].width = 40

    output.seek(0)
    return output.read()


def generate_pdf_report(df, title="Rapport", lang="fr"):
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    output = io.BytesIO()
    page_size = landscape(A4) if len(df.columns) > 6 else A4
    doc = SimpleDocTemplate(output, pagesize=page_size,
                             rightMargin=1.5*cm, leftMargin=1.5*cm,
                             topMargin=2*cm, bottomMargin=1.5*cm)

    C_RED = colors.HexColor("#E94560")
    C_DARK = colors.HexColor("#1A1A2E")
    C_NAVY = colors.HexColor("#16213E")
    C_WHITE = colors.white
    C_GRAY = colors.HexColor("#8A9BB4")
    C_LIGHT = colors.HexColor("#EAEAEA")

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("title", fontSize=22, textColor=C_WHITE,
                                  fontName="Helvetica-Bold", alignment=TA_CENTER, spaceAfter=6)
    subtitle_style = ParagraphStyle("subtitle", fontSize=10, textColor=C_GRAY,
                                     fontName="Helvetica", alignment=TA_CENTER, spaceAfter=20)
    section_style = ParagraphStyle("section", fontSize=13, textColor=C_RED,
                                    fontName="Helvetica-Bold", spaceBefore=15, spaceAfter=8)

    story = []
    story.append(Paragraph(title, title_style))
    date_str = datetime.now().strftime("%d/%m/%Y à %H:%M") if lang == "fr" else datetime.now().strftime("%m/%d/%Y at %H:%M")
    subtitle_text = f"Généré par DocFlow le {date_str}" if lang == "fr" else f"Generated by DocFlow on {date_str}"
    story.append(Paragraph(subtitle_text, subtitle_style))
    story.append(HRFlowable(width="100%", thickness=2, color=C_RED, spaceAfter=15))

    section_title = "Données" if lang == "fr" else "Data"
    story.append(Paragraph(section_title, section_style))

    max_cols = 8
    df_display = df.iloc[:500, :max_cols] if len(df.columns) > max_cols else df.iloc[:500]

    table_data = [df_display.columns.tolist()]
    for _, row in df_display.iterrows():
        table_data.append([str(v)[:30] if v is not None else "" for v in row.tolist()])

    page_width = landscape(A4)[0] - 3*cm if len(df.columns) > 6 else A4[0] - 3*cm
    col_width = page_width / len(df_display.columns)

    table = Table(table_data, colWidths=[col_width] * len(df_display.columns), repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), C_RED),
        ("TEXTCOLOR", (0, 0), (-1, 0), C_WHITE),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, 0), 9),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [C_NAVY, C_DARK]),
        ("TEXTCOLOR", (0, 1), (-1, -1), C_LIGHT),
        ("FONTSIZE", (0, 1), (-1, -1), 8),
        ("ROWHEIGHT", (0, 0), (-1, -1), 18),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#2A2A4A")),
    ]))
    story.append(table)

    numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
    if numeric_cols:
        story.append(Spacer(1, 15))
        stats_title = "Statistiques descriptives" if lang == "fr" else "Descriptive Statistics"
        story.append(Paragraph(stats_title, section_style))
        desc = df[numeric_cols].describe().round(2)
        stats_data = [[""] + desc.columns.tolist()]
        for idx_name, row in desc.iterrows():
            stats_data.append([str(idx_name)] + [str(v) for v in row.tolist()])
        stats_col_width = page_width / len(stats_data[0])
        stats_table = Table(stats_data, colWidths=[stats_col_width] * len(stats_data[0]), repeatRows=1)
        stats_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), C_RED),
            ("TEXTCOLOR", (0, 0), (-1, 0), C_WHITE),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [C_NAVY, C_DARK]),
            ("TEXTCOLOR", (0, 1), (-1, -1), C_LIGHT),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#2A2A4A")),
        ]))
        story.append(stats_table)

    doc.build(story)
    output.seek(0)
    return output.read()


def generate_word_report(df, title="Rapport", lang="fr"):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(11.7)
    section.page_height = Inches(8.3)

    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0xE9, 0x45, 0x60)
    run.font.size = Pt(24)

    date_str = datetime.now().strftime("%d/%m/%Y à %H:%M") if lang == "fr" else datetime.now().strftime("%m/%d/%Y")
    subtitle = f"Généré par DocFlow — {date_str}" if lang == "fr" else f"Generated by DocFlow — {date_str}"
    p = doc.add_paragraph(subtitle)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = RGBColor(0x8A, 0x9B, 0xB4)
    p.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    section_label = "Données" if lang == "fr" else "Data"
    h = doc.add_heading(section_label, level=1)
    h.runs[0].font.color.rgb = RGBColor(0xE9, 0x45, 0x60)

    df_display = df.iloc[:100, :10] if len(df.columns) > 10 else df.iloc[:100]

    table = doc.add_table(rows=1, cols=len(df_display.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df_display.columns):
        cell = hdr_cells[i]
        cell.text = str(col_name)
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "E94560")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:val"), "clear")
        tc_pr.append(shd)

    for row_idx, (_, row) in enumerate(df_display.iterrows()):
        cells = table.add_row().cells
        bg_color = "16213E" if row_idx % 2 == 0 else "1A1A2E"
        for col_idx, val in enumerate(row):
            cell = cells[col_idx]
            cell.text = str(val)[:40] if val is not None else ""
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(cell.text)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xEA, 0xEA, 0xEA)

    numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
    if numeric_cols:
        doc.add_paragraph()
        stats_label = "Statistiques descriptives" if lang == "fr" else "Descriptive Statistics"
        h2 = doc.add_heading(stats_label, level=1)
        h2.runs[0].font.color.rgb = RGBColor(0xE9, 0x45, 0x60)
        desc = df[numeric_cols].describe().round(2)
        stats_table = doc.add_table(rows=1, cols=len(desc.columns) + 1)
        stats_table.style = "Table Grid"
        hdr = stats_table.rows[0].cells
        hdr[0].text = ""
        for i, col in enumerate(desc.columns):
            hdr[i+1].text = str(col)
            hdr[i+1].paragraphs[0].runs[0].font.bold = True
        for idx_name, row in desc.iterrows():
            cells = stats_table.add_row().cells
            cells[0].text = str(idx_name)
            for i, val in enumerate(row):
                cells[i+1].text = str(val)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()
